# -*- coding: utf-8 -*-
"""
Gerador do Relatório de Probabilidade e Estatística — Star Dataset
Disciplina: Probabilidade e Estatística A — UFG

Uso:
    python scripts/gerar_relatorio.py
"""

import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Caminhos ──────────────────────────────────────────────────────────────────
BASE_DIR  = os.path.normpath(os.path.join(os.path.dirname(__file__), '..'))
DATA_FILE = os.path.join(BASE_DIR, '6 class csv.csv')
DOCX_OUT  = os.path.join(BASE_DIR, 'Relatório de Probabilidade e Estatística.docx')

IMG_QUAL = os.path.join(BASE_DIR, 'graficos', 'qualitativas')
IMG_DISC = os.path.join(BASE_DIR, 'graficos', 'discreta')
IMG_CONT = os.path.join(BASE_DIR, 'graficos', 'continuas')
IMG_ROOT = os.path.join(BASE_DIR, 'graficos')
IMG_REG  = os.path.join(BASE_DIR, 'graficos', 'regressao')

# ── Constantes de conteúdo ────────────────────────────────────────────────────
IMG_WIDTH = Inches(5.5)
FONT_NAME = 'Times New Roman'

STAR_TYPE_NAMES = {
    0: 'Anã Marrom',
    1: 'Anã Vermelha',
    2: 'Anã Branca',
    3: 'Seq. Principal',
    4: 'Supergigante',
    5: 'Hipergigante',
}

TEMP_BINS = [0, 3500, 5200, 6000, 7500, 30000, float('inf')]
TEMP_LABELS = [
    'Muito Fria (<3500 K)',
    'Fria (3500-5200 K)',
    'Morna (5200-6000 K)',
    'Moderada (6000-7500 K)',
    'Quente (7500-30000 K)',
    'Muito Quente (>30000 K)',
]
ORDEM_SPECTRAL = ['O', 'B', 'A', 'F', 'G', 'K', 'M']

# Contadores de figuras e tabelas (lista mutável para ser incrementável dentro de funções)
_fig_counter = [0]
_tab_counter = [0]


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE FORMATAÇÃO
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_font(cell, size=Pt(10), bold=False):
    """Aplica Times New Roman em todos os runs de uma célula de tabela."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = FONT_NAME
            run.font.size = size
            run.font.bold = bold


def _shade_cell(cell, fill='BFBFBF'):
    """Pinta o fundo de uma célula com a cor hex especificada."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)


def configure_document(doc):
    """Define margens ABNT, fonte e espaçamento padrão do documento."""
    sec = doc.sections[0]
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2)
    sec.top_margin    = Cm(3)
    sec.bottom_margin = Cm(2)

    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.5
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(6)


def add_paragraph(doc, text='', bold=False, italic=False,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, size=Pt(12)):
    """Adiciona parágrafo com fonte e espaçamento padronizados."""
    p = doc.add_paragraph(style='Normal')
    p.alignment = alignment
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing      = 1.5
    if text:
        run = p.add_run(text)
        run.font.name   = FONT_NAME
        run.font.size   = size
        run.font.bold   = bold
        run.font.italic = italic
    return p


def add_heading(doc, text, level=1):
    """
    Título de seção (level 1 = 14pt, level 2 = 12pt), negrito, Times New Roman.
    Não usa os estilos built-in Heading (que usam Calibri).
    """
    size = Pt(14) if level == 1 else Pt(12)
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before      = Pt(12)
    p.paragraph_format.space_after       = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing      = 1.5
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = True
    return p


def add_image(doc, img_path, caption=None):
    """Insere imagem centralizada com legenda numerada abaixo."""
    if not os.path.exists(img_path):
        add_paragraph(doc,
            f'[IMAGEM NAO ENCONTRADA: {os.path.basename(img_path)}]',
            italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return
    _fig_counter[0] += 1
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.add_run().add_picture(img_path, width=IMG_WIDTH)
    if caption:
        cap = doc.add_paragraph(style='Normal')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        r = cap.add_run(f'Figura {_fig_counter[0]}: {caption}')
        r.font.name   = FONT_NAME
        r.font.size   = Pt(10)
        r.font.italic = True


def add_word_table(doc, headers, rows, caption=None):
    """
    Cria tabela Word com cabeçalho cinza, fonte Times New Roman 10pt.
    Retorna o objeto Table.
    """
    _tab_counter[0] += 1
    if caption:
        cap = doc.add_paragraph(style='Normal')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(2)
        r = cap.add_run(f'Tabela {_tab_counter[0]}: {caption}')
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        r.font.bold = True

    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style    = 'Table Grid'
    table.autofit  = False

    # Largura igual para todas as colunas dentro da margem disponível
    col_w = Inches(6.0) / n_cols
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_w

    # Cabeçalho
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = hdr
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_font(cell, bold=True)
        _shade_cell(cell)

    # Dados
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_font(cell)

    # Espaço após a tabela
    doc.add_paragraph(style='Normal').paragraph_format.space_after = Pt(4)
    return table


def add_descriptive_measures_table(doc, d, var_name, unit=''):
    """Insere tabela 2-colunas com as 14 medidas descritivas calculadas."""
    headers = ['Medida Descritiva', f'Valor  ({unit})' if unit else 'Valor']
    rows = [
        ('Media',                  f"{d['mean']:.4f}"),
        ('Mediana',                f"{d['median']:.4f}"),
        ('Moda',                   f"{d['mode']:.4f}"),
        ('Desvio Padrao',          f"{d['std']:.4f}"),
        ('Variancia',              f"{d['var']:.4f}"),
        ('Coef. de Variacao (%)',  f"{d['cv']:.2f}"),
        ('Minimo',                 f"{d['min']:.4f}"),
        ('Maximo',                 f"{d['max']:.4f}"),
        ('Amplitude',              f"{d['range']:.4f}"),
        ('Q1 (25%)',               f"{d['q1']:.4f}"),
        ('Q3 (75%)',               f"{d['q3']:.4f}"),
        ('IIQ (Q3 - Q1)',          f"{d['iqr']:.4f}"),
        ('Assimetria de Pearson',  f"{d['skewness']:.4f}"),
        ('Curtose (Fisher)',        f"{d['kurtosis']:.4f}"),
    ]
    add_word_table(doc, headers, rows,
                   caption=f'Medidas Descritivas — {var_name}')


# ══════════════════════════════════════════════════════════════════════════════
# FUNÇÕES DE DADOS
# ══════════════════════════════════════════════════════════════════════════════

def _padronizar_cor(cor):
    c = cor.strip().lower()
    if c in ('blue',):
        return 'Blue'
    if c in ('blue white', 'blue-white'):
        return 'Blue White'
    if c in ('red',):
        return 'Red'
    if c in ('yellow white', 'yellow-white', 'yellowish white',
             'pale yellow orange', 'yellow'):
        return 'Yellow/White'
    if c in ('white', 'whitish'):
        return 'White'
    if c in ('orange', 'orange-red'):
        return 'Orange/Red'
    return cor.strip().title()


def load_data():
    df = pd.read_csv(DATA_FILE)
    df['Star type label'] = df['Star type'].map(STAR_TYPE_NAMES)
    df['Star color std']  = df['Star color'].apply(_padronizar_cor)
    df['Luminosity(L/Lo)'] = df['Luminosity(L/Lo)'].replace(0, np.nan)
    df['Radius(R/Ro)']     = df['Radius(R/Ro)'].replace(0, np.nan)
    df['log_Luminosity']   = np.log10(df['Luminosity(L/Lo)'])
    df['log_Radius']       = np.log10(df['Radius(R/Ro)'])
    df['Temp Categoria']   = pd.cut(
        df['Temperature (K)'], bins=TEMP_BINS, labels=TEMP_LABELS, right=False
    )
    return df


def build_freq_table_qualitative(series, order=None):
    contagem = series.value_counts()
    if order:
        contagem = contagem.reindex([o for o in order if o in contagem.index],
                                    fill_value=0)
    total = len(series)
    headers = ['Categoria', 'Freq. Absoluta', 'Freq. Relativa (%)',
               'Freq. Abs. Acumulada', 'Freq. Rel. Acumulada (%)']
    rows, acum_abs, acum_rel = [], 0, 0.0
    for cat, cnt in contagem.items():
        acum_abs += int(cnt)
        rel       = round(float(cnt) / total * 100, 2)
        acum_rel  = round(acum_rel + rel, 2)
        rows.append((cat, int(cnt), f'{rel:.2f}', acum_abs, f'{acum_rel:.2f}'))
    return headers, rows


def build_freq_table_discrete(series, bins, labels):
    total    = len(series)
    cat      = pd.cut(series, bins=bins, labels=labels, right=False)
    freq_abs = cat.value_counts().reindex(labels, fill_value=0)
    headers  = ['Categoria', 'Freq. Absoluta', 'Freq. Relativa (%)',
                'Freq. Abs. Acumulada', 'Freq. Rel. Acumulada (%)']
    rows, acum_abs, acum_rel = [], 0, 0.0
    for lbl in labels:
        cnt      = int(freq_abs[lbl])
        acum_abs += cnt
        rel       = round(cnt / total * 100, 2)
        acum_rel  = round(acum_rel + rel, 2)
        rows.append((lbl, cnt, f'{rel:.2f}', acum_abs, f'{acum_rel:.2f}'))
    return headers, rows


def build_freq_table_continuous(series, n_classes=6):
    s         = series.dropna()
    total     = len(s)
    cnts, bds = np.histogram(s, bins=n_classes)
    largura   = bds[1:] - bds[:-1]
    densidade = cnts / (total * largura)
    freq_rel  = cnts / total * 100
    headers   = ['Intervalo', 'Freq. Absoluta', 'Freq. Relativa (%)',
                 'Freq. Abs. Acumulada', 'Freq. Rel. Acumulada (%)', 'Densidade']
    rows, acum_abs, acum_rel = [], 0, 0.0
    for i in range(len(cnts)):
        intv      = f'[{bds[i]:.3g}, {bds[i+1]:.3g})'
        cnt       = int(cnts[i])
        acum_abs += cnt
        rel       = round(freq_rel[i], 2)
        acum_rel  = round(acum_rel + rel, 2)
        rows.append((intv, cnt, f'{rel:.2f}', acum_abs,
                     f'{acum_rel:.2f}', f'{densidade[i]:.6f}'))
    return headers, rows


def compute_descriptive(series):
    s    = series.dropna()
    mean = s.mean()
    med  = s.median()
    mode = s.mode().iloc[0] if len(s.mode()) > 0 else np.nan
    std  = s.std()
    var  = s.var()
    cv   = (std / mean * 100) if mean != 0 else np.nan
    mn   = s.min()
    mx   = s.max()
    return dict(
        mean=mean, median=med, mode=mode, std=std, var=var,
        cv=cv, min=mn, max=mx, range=mx - mn,
        q1=s.quantile(0.25), q3=s.quantile(0.75),
        iqr=s.quantile(0.75) - s.quantile(0.25),
        skewness=3 * (mean - med) / std if std != 0 else np.nan,
        kurtosis=stats.kurtosis(s, fisher=True),
    )


# ══════════════════════════════════════════════════════════════════════════════
# SEÇÕES DO DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

def build_cover(doc):
    add_paragraph(doc, '')
    add_paragraph(doc, 'UNIVERSIDADE FEDERAL DE GOIAS', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, 'BACHARELADO EM INTELIGENCIA ARTIFICIAL', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, 'DISCIPLINA: PROBABILIDADE E ESTATISTICA A', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        add_paragraph(doc, '')
    add_paragraph(doc, 'ANTONIO HENRIQUE QUEIROZ VILLAR LOPES',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, 'JAMILY VIEIRA GONCALVES',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, 'JULIA PEREIRA SOUZA',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        add_paragraph(doc, '')
    add_paragraph(doc, 'TRABALHO DE ESTATISTICA DESCRITIVA', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, 'Analise de Estrelas', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        add_paragraph(doc, '')
    add_paragraph(doc, 'DOCENTE: MARCIO AUGUSTO F. RODRIGUES',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        add_paragraph(doc, '')
    add_paragraph(doc, 'Goiania', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '2026',    alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def build_introducao(doc):
    add_heading(doc, '1. INTRODUCAO')
    add_paragraph(doc,
        'A classificacao estelar, processo de categorizar estrelas com base em suas '
        'propriedades fisicas, e um dos problemas fundamentais da Astronomia. O sistema '
        'moderno de classificacao de Morgan-Keenan (MK) agrupa estrelas de acordo com a '
        'temperatura superficial, luminosidade, cor e magnitude absoluta, resultando em '
        'seis tipos principais: Ana Marrom, Ana Vermelha, Ana Branca, Sequencia Principal, '
        'Supergigante e Hipergigante.')
    add_paragraph(doc,
        'Este trabalho tem como objetivo documentar as tecnicas de analise descritiva, '
        'estudadas na disciplina de Probabilidade e Estatistica, aplicadas aos dados '
        'obtidos do Star Dataset to Predict Star Types, disponivel no Kaggle '
        '(https://www.kaggle.com/datasets/deepu1109/star-dataset), contendo 240 '
        'observacoes de estrelas. A partir dessas informacoes, busca-se caracterizar '
        'e resumir as propriedades do conjunto de dados, identificando padroes, '
        'assimetrias e diferencas entre os tipos estelares. Todos os calculos e '
        'graficos foram obtidos com auxilio computacional utilizando a linguagem Python '
        '(bibliotecas pandas, numpy, matplotlib, seaborn e scipy).')


def build_descricao_dados(doc):
    add_heading(doc, '2. DESCRICAO DOS DADOS')
    add_paragraph(doc,
        'O dataset contem 240 observacoes de estrelas de seis tipos distintos, com '
        '40 estrelas de cada tipo (dataset balanceado por construcao). O tamanho amostral '
        'e n = 240, com 7 variaveis no total: 4 quantitativas continuas, 1 quantitativa '
        'discreta e 2 qualitativas.')

    headers = ['Variavel', 'Descricao', 'Tipo', 'Unidade/Escala']
    rows = [
        ('Temperature (K)',       'Temperatura superficial',         'Quant. discreta*',     'Kelvin (K)'),
        ('Luminosity (L/Lo)',     'Luminosidade relativa ao Sol',    'Quant. continua',      'L_sol (adim.)'),
        ('Radius (R/Ro)',         'Raio relativo ao Sol',            'Quant. continua',      'R_sol (adim.)'),
        ('Absolute Magnitude',    'Magnitude absoluta (brilho)',      'Quant. continua',      'Escala Mv'),
        ('Star color',            'Cor observada da estrela',        'Qualit. nominal',      'Cat. de cor'),
        ('Spectral Class',        'Classe espectral MK',             'Qualit. ordinal',      'O, B, A, F, G, K, M'),
        ('Star type',             'Tipo estelar (0 a 5)',            'Qualit. ordinal**',    '0, 1, 2, 3, 4, 5'),
    ]
    add_word_table(doc, headers, rows, caption='Descricao das variaveis do dataset')

    add_paragraph(doc,
        '*Obs. 1: Temperature e tecnicamente continua, mas registrada como inteiro '
        'em Kelvin, sem casas decimais, o que permite trata-la como discreta.')
    add_paragraph(doc,
        '**Obs. 2: Star_type representa uma classificacao ordinal: 0 = Ana Marrom; '
        '1 = Ana Vermelha; 2 = Ana Branca; 3 = Sequencia Principal; '
        '4 = Supergigante; 5 = Hipergigante.')


def build_qualitativas(doc, df):
    add_heading(doc, '3. VARIAVEIS QUALITATIVAS')
    add_paragraph(doc,
        'O dataset possui tres variaveis qualitativas: a cor da estrela (nominal), '
        'a classe espectral (ordinal, segundo a sequencia MK) e o tipo estelar '
        '(ordinal). Para variaveis qualitativas, as ferramentas estatisticas apropriadas '
        'sao as tabelas de frequencia, os graficos de barras e os graficos de setores.')

    # ── 3.1 Star Color ──────────────────────────────────────────────────────────
    add_heading(doc, '3.1 Cor da Estrela (Star Color)', level=2)
    add_paragraph(doc,
        'A variavel Star color e qualitativa nominal e registra a cor observada de cada '
        'estrela. O dataset original apresenta diversas variantes de grafia para a mesma '
        'cor (ex.: "Blue White", "Blue-white", "blue white"). As categorias foram '
        'padronizadas antes da analise.')

    hdrs, rows = build_freq_table_qualitative(df['Star color std'])
    add_word_table(doc, hdrs, rows, caption='Frequencias de Cor da Estrela (padronizado)')

    add_paragraph(doc,
        'A cor vermelha e dominante (46,67%), reflexo da grande quantidade de anas '
        'vermelhas e supergigantes vermelhas no dataset. As estrelas azuis e '
        'azul-brancas somam aproximadamente 40%, representando as estrelas mais quentes '
        'e massivas. As demais categorias (branca, amarela/branca, laranja) totalizam '
        'menos de 15% das observacoes.')

    add_image(doc, os.path.join(IMG_QUAL, 'barras_star_color.png'),
              caption='Grafico de Barras — Distribuicao de Cor da Estrela')
    add_image(doc, os.path.join(IMG_QUAL, 'pizza_star_color.png'),
              caption='Grafico de Pizza — Proporcao das Cores Estelares (Top 5 + Outros)')
    add_paragraph(doc,
        'O grafico de pizza confirma a predominancia da cor vermelha (46,7%), '
        'seguida pela azul (23,3%) e pela azul-branca (17,1%). As demais cores, '
        'agrupadas em "Outros", representam menos de 13% do total.')

    # ── 3.2 Spectral Class ──────────────────────────────────────────────────────
    add_heading(doc, '3.2 Classe Espectral (Spectral Class)', level=2)
    add_paragraph(doc,
        'A classe espectral e uma variavel qualitativa ordinal que segue a sequencia '
        'de Morgan-Keenan: O, B, A, F, G, K, M, em ordem decrescente de temperatura '
        'superficial.')

    hdrs, rows = build_freq_table_qualitative(df['Spectral Class'], order=ORDEM_SPECTRAL)
    add_word_table(doc, hdrs, rows, caption='Frequencias de Classe Espectral (sequencia O-M)')

    add_paragraph(doc,
        'A classe M (estrelas mais frias) e a mais frequente com 46,25%, incluindo '
        'anas vermelhas, anas marrons e supergigantes vermelhas. A classe B ocupa o '
        'segundo lugar (19,17%), enquanto a classe G (que inclui o Sol) e a mais rara '
        'no dataset, com apenas 1 estrela (0,42%).')

    add_image(doc, os.path.join(IMG_QUAL, 'barras_spectral_class.png'),
              caption='Grafico de Barras — Distribuicao das Classes Espectrais')

    # ── 3.3 Star Type ───────────────────────────────────────────────────────────
    add_heading(doc, '3.3 Tipo Estelar (Star Type)', level=2)
    add_paragraph(doc,
        'O tipo estelar e uma variavel qualitativa ordinal com seis categorias, '
        'ordenadas por porte/luminosidade crescente: Ana Marrom (0), Ana Vermelha (1), '
        'Ana Branca (2), Sequencia Principal (3), Supergigante (4) e Hipergigante (5).')

    hdrs, rows = build_freq_table_qualitative(
        df['Star type label'], order=list(STAR_TYPE_NAMES.values())
    )
    add_word_table(doc, hdrs, rows, caption='Frequencias de Tipo Estelar')

    add_paragraph(doc,
        'O dataset e perfeitamente balanceado: cada tipo estelar possui exatamente '
        '40 observacoes (16,67% do total). Esse balanceamento foi intencional na '
        'construcao do dataset original, garantindo representacao equitativa de todas '
        'as categorias.')

    add_image(doc, os.path.join(IMG_QUAL, 'barras_star_type.png'),
              caption='Grafico de Barras — Distribuicao dos Tipos Estelares')


def build_discreta(doc, df):
    add_heading(doc, '4. VARIAVEL QUANTITATIVA DISCRETA — TEMPERATURA (K)')
    add_paragraph(doc,
        'A temperatura superficial das estrelas e medida em Kelvin e registrada como '
        'valor inteiro no dataset, sem casas decimais. Por isso, e tratada como variavel '
        'quantitativa discreta. Os valores variam de 1.939 K (ana marrom muito fria) a '
        '40.000 K (hipergigante azul). Para facilitar a analise, os dados foram '
        'categorizados em 6 faixas de temperatura inspiradas nos limites da classificacao '
        'espectral de Morgan-Keenan.')

    # 4.1 Tabela de frequencias
    add_heading(doc, '4.1 Tabela de Frequencias', level=2)
    hdrs, rows = build_freq_table_discrete(
        df['Temperature (K)'], TEMP_BINS, TEMP_LABELS
    )
    add_word_table(doc, hdrs, rows,
                   caption='Frequencias de Temperatura por Faixa Espectral')
    add_paragraph(doc,
        'As faixas "Muito Fria" e "Quente" concentram a maior parte das estrelas '
        '(31,25% e 40,00%, respectivamente), refletindo os dois grupos dominantes do '
        'dataset: anas marrons/vermelhas e estrelas do tipo B/A. A faixa "Morna" '
        '(5200-6000 K), tipica de estrelas como o Sol, e a menos representada '
        '(apenas 5 estrelas, 2,08%).')

    # 4.2 Histograma e poligono
    add_heading(doc, '4.2 Histograma e Poligono de Frequencias', level=2)
    add_image(doc, os.path.join(IMG_DISC, 'histograma_poligono_temperatura.png'),
              caption='Histograma e Poligono de Frequencias — Temperatura (K)')
    add_paragraph(doc,
        'O histograma revela uma distribuicao bimodal, com picos na faixa de baixas '
        'temperaturas (2000-5000 K) e na faixa de altas temperaturas (7500-30000 K), '
        'refletindo os dois grandes grupos de estrelas. O poligono de frequencias '
        'confirma essa bimodalidade, com uma regiao de baixa frequencia entre 5000 '
        'e 7500 K separando os dois picos.')

    # 4.3 Medidas descritivas
    add_heading(doc, '4.3 Medidas Descritivas', level=2)
    d = compute_descriptive(df['Temperature (K)'])
    add_descriptive_measures_table(doc, d, 'Temperature (K)', unit='K')
    add_paragraph(doc,
        f'A temperatura media ({d["mean"]:.0f} K) e muito superior a mediana '
        f'({d["median"]:.0f} K), indicando forte assimetria positiva: poucas estrelas '
        f'muito quentes elevam a media. A moda ({d["mode"]:.0f} K) corresponde a '
        f'estrelas da classe M (anas vermelhas e anas marrons). '
        f'O coeficiente de variacao de {d["cv"]:.1f}% confirma a heterogeneidade '
        f'extrema da amostra. '
        f'A assimetria de Pearson ({d["skewness"]:.2f}) confirma a cauda direita '
        f'(assimetria positiva), e a curtose ({d["kurtosis"]:.2f}) indica '
        f'distribuicao leptocurtica (caudas mais pesadas que a normal).')

    # 4.4 Boxplot
    add_heading(doc, '4.4 Boxplot', level=2)
    add_image(doc, os.path.join(IMG_DISC, 'boxplot_temperatura.png'),
              caption='Boxplot Geral e por Tipo Estelar — Temperatura (K)')
    add_paragraph(doc,
        f'O boxplot geral mostra uma caixa estreita (IIQ = {d["iqr"]:.0f} K) '
        f'em relacao a amplitude total ({d["range"]:.0f} K), com outliers no limite '
        f'superior correspondentes a estrelas hipergigantes e supergigantes de '
        f'temperatura muito elevada. O boxplot por tipo estelar revela que anas '
        f'marrons e vermelhas concentram-se em temperaturas muito baixas '
        f'(mediana < 3.500 K), enquanto hipergigantes e a sequencia principal '
        f'apresentam temperaturas mais elevadas com maior dispersao interna.')


def _continuous_section(doc, df, col, titulo, unit, hist_file, box_file,
                         interp_tab, interp_hist, interp_box):
    """Bloco padrao para uma sub-secao de variavel continua."""
    add_heading(doc, titulo, level=2)

    hdrs, rows = build_freq_table_continuous(df[col])
    add_word_table(doc, hdrs, rows,
                   caption=f'Frequencias com Densidade — {titulo}')
    add_paragraph(doc, interp_tab)

    add_image(doc, os.path.join(IMG_CONT, hist_file),
              caption=f'Histograma com KDE — {titulo}')
    add_paragraph(doc, interp_hist)

    add_image(doc, os.path.join(IMG_CONT, box_file),
              caption=f'Boxplot por Tipo Estelar — {titulo}')
    add_paragraph(doc, interp_box)

    d = compute_descriptive(df[col])
    add_descriptive_measures_table(doc, d, titulo, unit=unit)
    return d


def build_continuas(doc, df):
    add_heading(doc, '5. VARIAVEIS QUANTITATIVAS CONTINUAS')
    add_paragraph(doc,
        'As variaveis continuas do dataset sao: Luminosidade (L/Lo), Raio (R/Ro) e '
        'Magnitude Absoluta (Mv). Luminosidade e Raio apresentam distribuicoes '
        'fortemente assimetricas na escala original — hipergigantes sao ordens de '
        'magnitude maiores que anas marrons — motivando a analise tambem em escala '
        'logaritmica (log10). A magnitude absoluta ja esta em escala logaritmica '
        'por definicao astronomica.')

    # 5.1 Luminosity original
    d = _continuous_section(
        doc, df, 'Luminosity(L/Lo)',
        titulo='5.1 Luminosidade — escala original',
        unit='L/Lo',
        hist_file='hist_luminosity.png',
        box_file='boxplot_luminosity.png',
        interp_tab=(
            'A tabela evidencia a assimetria extrema: a primeira classe concentra '
            'quase a totalidade das observacoes, enquanto hipergigantes (L > 10^5 L_sol) '
            'pertencem a classes praticamente vazias. Essa concentracao torna a escala '
            'original inadequada para visualizacao direta.'),
        interp_hist=(
            'O histograma confirma a distribuicao extremamente assimetrica a direita. '
            'A curva KDE mostra densidade praticamente nula para luminosidades '
            'intermediarias, com pico estreito proximo de zero e cauda muito longa '
            'para valores altos.'),
        interp_box=(
            'O boxplot por tipo estelar revela que anas marrons, vermelhas e brancas '
            'possuem luminosidade tao baixa que suas caixas ficam comprimidas junto ao '
            'eixo. Hipergigantes e supergigantes aparecem como outliers extremos, '
            'demonstrando a necessidade da transformacao logaritmica.'),
    )

    # 5.2 log10(Luminosity)
    d2 = _continuous_section(
        doc, df, 'log_Luminosity',
        titulo='5.2 Luminosidade — escala log10',
        unit='log10(L/Lo)',
        hist_file='hist_log_luminosity.png',
        box_file='boxplot_log_luminosity.png',
        interp_tab=(
            'Em escala logaritmica a distribuicao se torna muito mais simetrica, com '
            'frequencias distribuidas de forma mais uniforme entre as classes. '
            'Isso confirma que a luminosidade estelar segue aproximadamente uma '
            'distribuicao log-normal.'),
        interp_hist=(
            'O histograma em log10 apresenta estrutura bimodal, com um pico para '
            'estrelas compactas (log L < -2) e outro para estrelas gigantes (log L > 3). '
            'A curva KDE revela essa estrutura com clareza.'),
        interp_box=(
            'O boxplot por tipo estelar em escala log mostra separacao clara entre '
            'as categorias. Anas marrons e brancas concentram-se em luminosidades '
            'baixas, enquanto hipergigantes ocupam o extremo superior. A sequencia '
            'principal apresenta a maior variabilidade interna.'),
    )
    add_paragraph(doc,
        f'Apos a transformacao logaritmica, a assimetria de Pearson reduz '
        f'significativamente (de {d["skewness"]:.2f} para {d2["skewness"]:.2f}), '
        f'e o CV cai de {d["cv"]:.1f}% para {d2["cv"]:.1f}%, confirmando que '
        f'a escala log lineariza a distribuicao.')

    # 5.3 Radius original
    d3 = _continuous_section(
        doc, df, 'Radius(R/Ro)',
        titulo='5.3 Raio — escala original',
        unit='R/Ro',
        hist_file='hist_radius.png',
        box_file='boxplot_radius.png',
        interp_tab=(
            'Assim como a luminosidade, o raio estelar apresenta assimetria extrema '
            'na escala original. A primeira classe concentra a maioria das estrelas '
            '(anas com raio < 1 R_sol), enquanto hipergigantes com raios acima de '
            '50 R_sol formam a cauda direita.'),
        interp_hist=(
            'O histograma confirma a forte assimetria positiva: a densidade e maxima '
            'proxima de zero e decresce rapidamente. A cauda direita e muito longa, '
            'representando as estrelas gigantes.'),
        interp_box=(
            'O boxplot por tipo estelar mostra que anas brancas possuem raios '
            'minusculos (< 0,02 R_sol), enquanto supergigantes e hipergigantes '
            'apresentam raios de dezenas a centenas de vezes o raio solar.'),
    )

    # 5.4 log10(Radius)
    d4 = _continuous_section(
        doc, df, 'log_Radius',
        titulo='5.4 Raio — escala log10',
        unit='log10(R/Ro)',
        hist_file='hist_log_radius.png',
        box_file='boxplot_log_radius.png',
        interp_tab=(
            'A transformacao logaritmica distribui as classes de forma mais equilibrada. '
            'A estrutura bimodal do log(Raio) reflete dois grupos naturais: estrelas '
            'compactas (anas) com log R < 0 e estrelas gigantes com log R > 1.'),
        interp_hist=(
            'O histograma em log10 revela dois picos bem definidos: um para estrelas '
            'compactas (log R proximo de -2 a -1) e outro para estrelas de grande '
            'porte (log R proximo de 1 a 2). A curva KDE evidencia essa bimodalidade, '
            'refletindo as duas populacoes fisicas distintas do dataset.'),
        interp_box=(
            'O boxplot por tipo estelar em escala log confirma a separacao clara entre '
            'os grupos. Anas brancas tem os menores raios logaritmicos, enquanto '
            'hipergigantes apresentam os maiores. A sequencia principal abrange '
            'uma ampla faixa de raios, refletindo sua heterogeneidade interna.'),
    )
    add_paragraph(doc,
        f'A transformacao log do Raio tambem reduz a assimetria de Pearson '
        f'(de {d3["skewness"]:.2f} para {d4["skewness"]:.2f}) e o CV '
        f'(de {d3["cv"]:.1f}% para {d4["cv"]:.1f}%), confirmando o comportamento '
        f'log-normal do raio estelar.')

    # 5.5 Magnitude Absoluta
    _continuous_section(
        doc, df, 'Absolute magnitude(Mv)',
        titulo='5.5 Magnitude Absoluta (Mv)',
        unit='Mv',
        hist_file='hist_magnitude.png',
        box_file='boxplot_magnitude.png',
        interp_tab=(
            'A magnitude absoluta varia de -11,92 (hipergigantes brilhantissimas) '
            'a +20,06 (anas marrons muito fracas). A escala e invertida: valores mais '
            'negativos correspondem a estrelas mais brilhantes. A distribuicao '
            'apresenta frequencias relativamente distribuidas entre as classes.'),
        interp_hist=(
            'O histograma revela uma distribuicao bimodal, com picos em torno de '
            'Mv ~ -6 (supergigantes/hipergigantes) e Mv ~ +12 (anas marrons/vermelhas). '
            'A regiao central e menos frequente. A curtose negativa indica '
            'distribuicao platicurtica (caudas mais leves que a normal).'),
        interp_box=(
            'O boxplot por tipo estelar mostra separacao muito clara entre categorias: '
            'hipergigantes e supergigantes tem magnitudes muito negativas (muito '
            'brilhantes), enquanto anas marrons apresentam magnitudes muito positivas '
            '(muito fracas). A sequencia principal e anas brancas ocupam a faixa '
            'intermediaria.'),
    )


def build_diagrama_hr(doc):
    add_heading(doc, '6. DIAGRAMA DE HERTZSPRUNG-RUSSELL')
    add_paragraph(doc,
        'O Diagrama de Hertzsprung-Russell (HR) e uma das ferramentas mais importantes '
        'da Astronomia estelar. Ele plota a luminosidade de uma estrela (eixo Y, em '
        'escala log10) em funcao de sua temperatura superficial (eixo X, em escala '
        'logaritmica e invertida — temperatura decresce da esquerda para a direita, '
        'por convencao astronomica). Cada ponto representa uma estrela, colorido '
        'pelo tipo estelar e com tamanho proporcional ao raio estelar.')

    add_image(doc, os.path.join(IMG_ROOT, 'diagrama_hr.png'),
              caption='Diagrama de Hertzsprung-Russell — log10(Luminosidade) x Temperatura (K)')

    add_paragraph(doc,
        'O diagrama reproduz com clareza as estruturas caracteristicas esperadas pela '
        'teoria astronomica:')
    add_paragraph(doc,
        'Sequencia Principal (tipo 3): forma uma faixa diagonal do canto superior '
        'esquerdo (estrelas quentes e brilhantes) ao inferior direito (estrelas frias '
        'e fracas). E onde a maioria das estrelas passa a maior parte de sua vida.')
    add_paragraph(doc,
        'Anas Brancas (tipo 2): isoladas no canto inferior esquerdo — alta temperatura '
        'mas baixissima luminosidade, refletindo seu pequeno tamanho (raio ~ 0,01 R_sol).')
    add_paragraph(doc,
        'Anas Marrons e Vermelhas (tipos 0 e 1): agrupadas no canto inferior direito, '
        'com baixas temperaturas e baixissimas luminosidades.')
    add_paragraph(doc,
        'Supergigantes e Hipergigantes (tipos 4 e 5): dispersas na parte superior do '
        'diagrama (alta luminosidade), tanto em temperaturas altas (azuis) quanto '
        'baixas (vermelhas). Os pontos maiores (raio maior) correspondem as hipergigantes.')


def build_regressao(doc, df):
    add_heading(doc, '7. CORRELACAO E REGRESSAO LINEAR SIMPLES')
    add_paragraph(doc,
        'A analise de correlacao e regressao linear permite quantificar e modelar '
        'as relacoes lineares entre as variaveis numericas do dataset. A correlacao '
        'de Pearson mede a forca e a direcao da relacao linear, variando de -1 '
        '(correlacao negativa perfeita) a +1 (correlacao positiva perfeita). '
        'A regressao linear simples ajusta uma reta y = b0 + b1*x para o par '
        'de maior correlacao absoluta identificado automaticamente.')

    # 7.1 Matriz de correlacao
    add_heading(doc, '7.1 Matriz de Correlacao de Pearson', level=2)
    add_paragraph(doc,
        'A matriz de correlacao foi calculada para todas as variaveis numericas, '
        'incluindo as versoes log10 de Luminosidade e Raio, pois estas apresentam '
        'relacoes mais lineares com as demais variaveis.')

    add_image(doc, os.path.join(IMG_REG, 'heatmap_correlacao.png'),
              caption='Heatmap da Matriz de Correlacao de Pearson')

    add_paragraph(doc,
        'As correlacoes mais expressivas encontradas sao:')
    add_paragraph(doc,
        '• log10(Luminosidade) x Magnitude Absoluta: r = -0,977 (correlacao negativa '
        'muito forte). O sinal negativo e esperado: maior luminosidade implica '
        'magnitude mais negativa (estrela mais brilhante).')
    add_paragraph(doc,
        '• log10(Raio) x Magnitude Absoluta: r = -0,912 (correlacao negativa forte). '
        'Estrelas maiores tendem a ser mais brilhantes.')
    add_paragraph(doc,
        '• log10(Luminosidade) x log10(Raio): r = 0,918 (correlacao positiva forte). '
        'Estrelas maiores tendem a ter maior luminosidade.')

    # 7.2 Regressao linear
    add_heading(doc, '7.2 Modelo de Regressao Linear Simples', level=2)
    add_paragraph(doc,
        'O par com maior correlacao absoluta identificado automaticamente foi '
        'log10(Luminosidade) x Magnitude Absoluta (r = -0,9766). '
        'O modelo de regressao linear ajustado e:')

    p_eq = doc.add_paragraph(style='Normal')
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_eq = p_eq.add_run('Mv  =  -2,5948 x log10(L/Lo)  +  6,2172')
    r_eq.font.name = FONT_NAME
    r_eq.font.size = Pt(12)
    r_eq.font.bold = True

    reg_headers = ['Estatistica', 'Valor']
    reg_rows = [
        ('Coeficiente angular (b1)', '-2,5948'),
        ('Coeficiente linear  (b0)',  '6,2172'),
        ('Correlacao de Pearson (r)', '-0,9766'),
        ('R2 (coef. de determinacao)', '95,37%'),
        ('RMSE', '2,2615'),
        ('p-valor', '8,33 x 10^-161'),
    ]
    add_word_table(doc, reg_headers, reg_rows,
                   caption='Resultados da Regressao Linear — log10(L/Lo) x Mv')

    add_paragraph(doc,
        'O R2 = 95,37% indica que 95,37% da variancia na magnitude absoluta e '
        'explicada linearmente pelo logaritmo da luminosidade — um ajuste excelente. '
        'O coeficiente angular de -2,5948 esta muito proximo do valor teorico de '
        '-2,5 previsto pela definicao da escala de magnitude astronomica '
        '(m1 - m2 = -2,5 x log10(L1/L2)), validando a consistencia interna do '
        'dataset. O p-valor extremamente pequeno confirma a significancia estatistica '
        'do modelo.')

    # 7.3 Dispersao
    add_heading(doc, '7.3 Grafico de Dispersao com Reta Ajustada', level=2)
    add_image(doc, os.path.join(IMG_REG, 'dispersao_regressao.png'),
              caption='Dispersao com Reta Ajustada — log10(L/Lo) x Magnitude Absoluta (Mv)')
    add_paragraph(doc,
        'O grafico de dispersao mostra os pontos concentrados muito proximos a reta '
        'ajustada ao longo de toda a amplitude dos dados, consistente com o alto R2. '
        'A equacao da reta esta anotada no proprio grafico. A reta cobre desde as '
        'anas marrons (baixa luminosidade, alta magnitude) ate as hipergigantes '
        '(alta luminosidade, magnitude muito negativa).')

    # 7.4 Residuos
    add_heading(doc, '7.4 Analise de Residuos', level=2)
    add_image(doc, os.path.join(IMG_REG, 'residuos_regressao.png'),
              caption='Grafico de Residuos vs. Valores Ajustados')
    add_paragraph(doc,
        'Os residuos estao distribuidos de forma aproximadamente aleatoria ao redor '
        'de zero, sem padrao sistematico evidente — suportando a suposicao de '
        'homocedasticidade (variancia constante dos erros). Os poucos pontos com '
        'residuos maiores nas extremidades correspondem a estrelas de tipos '
        'menos comuns na faixa de transicao entre categorias, o que e esperado '
        'dado que o modelo e global (nao especifico por tipo estelar).')


def build_referencias(doc):
    add_heading(doc, 'REFERENCIAS')

    refs = [
        # Dataset
        'DEEPU1109. Star Dataset to Predict Star Types. Kaggle, 2021. '
        'Disponivel em: https://www.kaggle.com/datasets/deepu1109/star-dataset. '
        'Acesso em: jun. 2026.',

        # Estatistica descritiva (livro-texto principal da disciplina)
        'BUSSAB, Wilton O.; MORETTIN, Pedro A. Estatistica Basica. '
        '9. ed. Sao Paulo: Saraiva Educacao, 2017. 576 p. ISBN 978-85-472-0063-6.',

        # Livro de referencia para regressao e correlacao
        'MONTGOMERY, Douglas C.; RUNGER, George C. Estatistica Aplicada e '
        'Probabilidade para Engenheiros. 6. ed. Rio de Janeiro: LTC, 2016. 629 p. '
        'ISBN 978-85-216-3360-5.',

        # Astronomia: classificacao estelar e diagrama HR
        'CARROLL, Bradley W.; OSTLIE, Dale A. An Introduction to Modern Astrophysics. '
        '2. ed. Cambridge: Cambridge University Press, 2017. 1359 p. '
        'ISBN 978-1-108-42216-1.',

        # Sistema de classificacao espectral MK
        'MORGAN, William W.; KEENAN, Philip C.; KELLMAN, Edith. An Atlas of Stellar '
        'Spectra, with an Outline of Spectral Classification. Chicago: University of '
        'Chicago Press, 1943.',

        # NumPy
        'HARRIS, Charles R. et al. Array programming with NumPy. '
        'Nature, v. 585, n. 7825, p. 357-362, set. 2020. '
        'DOI: 10.1038/s41586-020-2649-2.',

        # pandas
        'MCKINNEY, Wes. Data Structures for Statistical Computing in Python. '
        'In: PYTHON IN SCIENCE CONFERENCE, 9., 2010, Austin. '
        'Proceedings... Austin: SciPy, 2010. p. 56-61. '
        'DOI: 10.25080/Majora-92bf1922-00a.',

        # matplotlib
        'HUNTER, John D. Matplotlib: A 2D Graphics Environment. '
        'Computing in Science & Engineering, v. 9, n. 3, p. 90-95, 2007. '
        'DOI: 10.1109/MCSE.2007.55.',

        # seaborn
        'WASKOM, Michael. seaborn: statistical data visualization. '
        'Journal of Open Source Software, v. 6, n. 60, p. 3021, 2021. '
        'DOI: 10.21105/joss.03021.',

        # SciPy
        'VIRTANEN, Pauli et al. SciPy 1.0: Fundamental Algorithms for Scientific '
        'Computing in Python. Nature Methods, v. 17, p. 261-272, 2020. '
        'DOI: 10.1038/s41592-019-0686-2.',

        # python-docx
        'PYTHON-DOCX CONTRIBUTORS. python-docx: Create and update Microsoft Word '
        '.docx files. Versao 1.2.0. GitHub, 2024. '
        'Disponivel em: https://github.com/python-openxml/python-docx. '
        'Acesso em: jun. 2026.',

        # Python
        'PYTHON SOFTWARE FOUNDATION. Python Language Reference. Versao 3.14. '
        'Wilmington, DE: PSF, 2025. '
        'Disponivel em: https://www.python.org. Acesso em: jun. 2026.',
    ]

    for ref in refs:
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after  = Pt(6)
        # Recuo ABNT: primeira linha alinhada, demais com recuo de 1,25 cm
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent       = Cm(0)
        run = p.add_run(ref)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)


def build_conclusao(doc):
    add_heading(doc, '8. CONCLUSAO')
    add_paragraph(doc,
        'Este trabalho realizou uma analise estatistica descritiva completa do '
        'Star Dataset, com 240 estrelas de seis tipos distintos. Foram analisadas '
        'tres variaveis qualitativas (cor, classe espectral e tipo estelar), uma '
        'variavel quantitativa discreta (temperatura) e tres continuas (luminosidade, '
        'raio e magnitude absoluta), alem da construcao do Diagrama de '
        'Hertzsprung-Russell e da analise de correlacao e regressao linear.')
    add_paragraph(doc,
        'Os resultados revelam caracteristicas fisicas fundamentais das estrelas: '
        'a variabilidade extrema das variaveis quantitativas (CV > 90% para '
        'temperatura na escala original), a natureza log-normal da luminosidade e '
        'do raio, e a bimodalidade das distribuicoes que refletem a separacao fisica '
        'entre estrelas compactas e gigantes. O Diagrama HR permitiu visualizar '
        'claramente as diferentes populacoes estelares, replicando o padrao classico '
        'da Astronomia com os dados do dataset.')
    add_paragraph(doc,
        'A regressao linear entre log10(Luminosidade) e Magnitude Absoluta produziu '
        'R2 = 95,37%, com coeficiente angular de -2,5948, muito proximo do valor '
        'teorico -2,5 da escala de magnitude astronomica. Esse resultado demonstra '
        'a consistencia interna dos dados e ilustra como a analise estatistica '
        'descritiva pode revelar relacoes fisicas profundas, mesmo sem recorrer a '
        'modelos preditivos complexos. Todos os resultados foram obtidos com '
        'auxilio computacional em linguagem Python.')


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    print('Carregando dados...')
    df = load_data()

    print('Criando documento Word...')
    doc = Document()
    configure_document(doc)

    build_cover(doc);           print('  [OK] Capa')
    build_introducao(doc);      print('  [OK] Introducao')
    build_descricao_dados(doc); print('  [OK] Descricao dos dados')
    build_qualitativas(doc, df);print('  [OK] Variaveis qualitativas')
    build_discreta(doc, df);    print('  [OK] Variavel discreta')
    build_continuas(doc, df);   print('  [OK] Variaveis continuas')
    build_diagrama_hr(doc);     print('  [OK] Diagrama HR')
    build_regressao(doc, df);   print('  [OK] Correlacao e regressao')
    build_conclusao(doc);       print('  [OK] Conclusao')
    build_referencias(doc);     print('  [OK] Referencias')

    doc.save(DOCX_OUT)
    print(f'\nDocumento salvo em: {DOCX_OUT}')
    print(f'Total de figuras:   {_fig_counter[0]}')
    print(f'Total de tabelas:   {_tab_counter[0]}')


if __name__ == '__main__':
    main()
